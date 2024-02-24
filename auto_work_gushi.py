import openpyxl
import pandas as pd
import re,os
from docx import Document 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime



def read_data():
    df = pd.read_excel('台账数据.xlsx' )
    #打印所有的项目序号及名称
    print(df.iloc[:, 0:2].to_string(index=False))
    row_num = int(input("请输入项目序号：")) # 让用户输入要读取的行数
    df_data = df.iloc[[row_num-1]]
    # print(df_data.to_string(index=False))
    # # 保存到csv文件
    # df_data.to_csv('data.csv')
    # data= pd.read_csv('data.csv' )
    return df_data

def make_decision_sheet(data):    #定案表  
    wb = openpyxl.load_workbook(r'.\template_gushi\4定案表模板.xlsx')
    sheet = wb['Sheet1']    
    #从data中读取数据到sheet中
    sheet.cell(row=2, column=2).value = data.iloc[0]['工程名称']
    sheet.cell(row=3, column=2).value = data.iloc[0]['设计单位']
    sheet.cell(row=4, column=2).value = data.iloc[0]['送审预算编制单位']
    sheet.cell(row=2, column=4).value = data.iloc[0]['送审金额']
    sheet.cell(row=3, column=4).value = data.iloc[0]['审定金额']
    sheet.cell(row=2, column=6).value = '委托书编号：\n\n'+data.iloc[0]['委托书编号']
    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out','附件2定案表-'+data.iloc[0]['工程名称']+'.xlsx')
    wb.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)


def make_govmark(data):    #处理签  
    wb = openpyxl.load_workbook(r'.\template_gushi\处理签模板.xlsx')
    sheet = wb['Sheet1'] 
    project_name=data.iloc[0]['工程名称']   
    sheet.cell(row=7, column=2).value = '关于'+project_name+'预算的评审意见'
    #保存的名称以变量"工程名称"命名
    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out','处理签-'+project_name+'.xlsx')
    wb.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)


#word中查找关键词替换
def docx_replace_regex(doc_obj, regex , replace):  
    for p in doc_obj.paragraphs: 
        if regex.search(p.text): 
            inline = p.runs 
            # Loop added to work with runs (strings with same style) 
            for i in range(len(inline)): 
                if regex.search(inline[i].text): 
                    text = regex.sub(str(replace), inline[i].text) 
                    inline[i].text = text 
 
    for table in doc_obj.tables: 
        for row in table.rows: 
            for cell in row.cells: 
                docx_replace_regex(cell, regex , replace) 
 
 
 
def make_Application_word(data):
    doc = Document(r'.\template_gushi\1_1初审复核申请模板.docx')
    for index, row in data.iterrows():
        regex1 = re.compile(r"工程名称") 
        regex2 = re.compile(r"送审的金额") 
        regex3 = re.compile(r"初审的金额") 
        regex4 = re.compile(r"审减的金额") 
        regex5 = re.compile(r"初审复核时间") 

        replace1 = row['工程名称']
        replace2 = row['送审金额']
        num2 = float(replace2 )
        replace2_num = "{:.2f}".format(num2)
        replace3 = row['提交初审金额']
        num3 = float(replace3)
        replace3_num = "{:.2f}".format(num3)
        replace4 = row['审减金额']
        num4= float(replace4)
        replace4_num= "{:.2f}".format(num4)
        replace5_date = row['初审复核时间']
        # replace5_date = datetime.datetime.strptime(replace5_str,"%Y-%m-%d") #字符串转日期
        replace5=f'{replace5_date.year}年{replace5_date.month}月{replace5_date.day}日' 

        docx_replace_regex(doc, regex1 , replace1) 
        docx_replace_regex(doc, regex2 , replace2_num)
        docx_replace_regex(doc, regex3 , replace3_num)
        docx_replace_regex(doc, regex4 , replace4_num)
        docx_replace_regex(doc, regex5 , replace5)
    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out','初审复核申请-'+row["工程名称"]+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)
    
def make_check_word(data):     #复核意见
    doc = Document(r'.\template_gushi\1_2项目复核意见模板.docx')
    for index, row in data.iterrows():
        regex1 = re.compile(r"工程名称") 
        regex3 = re.compile(r"初审的金额") 
        regex5 = re.compile(r"初审复核时间") 

        replace1 = row['工程名称']
        replace3 = row['提交初审金额']
        num3 = float(replace3)
        replace3_num = "{:.2f}".format(num3)
        replace5_date = row['初审复核时间']
        replace5=f'{replace5_date.year}年{replace5_date.month}月{replace5_date.day}日' 

        docx_replace_regex(doc, regex1 , replace1) 
        docx_replace_regex(doc, regex3 , replace3_num)
        docx_replace_regex(doc, regex5 , replace5)
    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out','复核意见-'+row["工程名称"]+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)

def make_check_reply_word(data):
    doc = Document(r'.\template_gushi\1_3复核意见回复模板.docx')
    for index, row in data.iterrows():
        regex1 = re.compile(r"工程名称") 
        regex2 = re.compile(r"送审的金额") 
        regex3 = re.compile(r"初审的金额") 
        regex3_1 = re.compile(r"审定的金额")
        regex3_2 = re.compile(r"审定减初审") 
        regex4 = re.compile(r"审减的金额") 
        regex5 = re.compile(r"初审复核时间") 

        replace1 = row['工程名称']
        replace2 = row['送审金额']
        num2 = float(replace2 )
        replace2_num = "{:.2f}".format(num2)
        replace3 = row['提交初审金额']
        num3 = float(replace3)
        replace3_num = "{:.2f}".format(num3)
        replace3_1 = row['审定金额']
        num3_1 = float(replace3_1)
        replace3_num_1 = "{:.2f}".format(num3_1)
        replace3_2 = replace3_1 - replace3 #审定金额-初审金额
        num3_2 = float(replace3_2)
        replace3_num_2 = "{:.2f}".format(num3_2)
        if  replace3_2 > 0:
            replace3_str_2 = '审增' + replace3_num_2
        else:
            replace3_str_2 = '审减' + replace3_num_2
        replace4 = row['审减金额']
        num4= float(replace4)
        replace4_num= "{:.2f}".format(num4)
        replace5_date = row['初审复核时间']
        replace5=f'{replace5_date.year}年{replace5_date.month}月{replace5_date.day}日' 

        docx_replace_regex(doc, regex1 , replace1) 
        docx_replace_regex(doc, regex2 , replace2_num)
        docx_replace_regex(doc, regex3 , replace3_num)
        docx_replace_regex(doc, regex3_1 , replace3_num_1)
        docx_replace_regex(doc, regex3_2 , replace3_str_2)
        docx_replace_regex(doc, regex4 , replace4_num)
        docx_replace_regex(doc, regex5 , replace5)
    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out','复核意见回复-'+row["工程名称"]+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)


def make_feedback_word(data):
    doc = Document(r'.\template_gushi\3反馈函模板.docx')
    for index, row in data.iterrows():
        regex1 = re.compile(r"工程名称") 
        regex2 = re.compile(r"送审金额") 
        regex3 = re.compile(r"审定金额") 
        regex4 = re.compile(r"审减金额") 
        regex5 = re.compile(r"反馈时间") 
        regex6 = re.compile(r"送审甲方")
        

        replace1 = row['工程名称']
        replace2 = row['送审金额']
        num2 = float(replace2 )
        replace2_num = "{:.2f}".format(num2)
        replace3 = row['审定金额']
        num3 = float(replace3)
        replace3_num = "{:.2f}".format(num3)
        replace4 = row['审减金额']
        num4= float(replace4)
        replace4_num= "{:.2f}".format(num4)
        replace5_date = row['反馈的时间']
        # replace5_date = datetime.datetime.strptime(replace5_str,"%Y-%m-%d") #字符串转日期
        replace5=f'{replace5_date.year}年{replace5_date.month}月{replace5_date.day}日'
        replace6 = row['送审甲方']

        docx_replace_regex(doc, regex1 , replace1) 
        docx_replace_regex(doc, regex2 , replace2_num)
        docx_replace_regex(doc, regex3 , replace3_num)
        docx_replace_regex(doc, regex4 , replace4_num)
        docx_replace_regex(doc, regex5 , replace5)
        docx_replace_regex(doc, regex6 , replace6)

    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out','反馈函-'+row["工程名称"]+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)
def make_deduction_details_word(data):#审减明细
    replacements = {
        r"评审工程名": "工程名称",
        r"送审报价": "送审金额",
        r"评审价": "审定金额",
        r"审减价": "审减金额",
        r"今天日期": "今天日期",
        r"送审甲方": "送审甲方",
        r"设计公司": "设计单位",
        r"送审预算单位": "送审预算编制单位",
        r"评审范围内容" : "评审范围内容",
    }      
    doc = Document(r'.\template_gushi\2评审结果说明模板.docx')  
    for index, row in data.iterrows():          
        for regex, replace in replacements.items():          
            regex = re.compile(regex, re.DOTALL)
            if replace ==  "送审金额"  or replace == "审定金额" or replace == "审减金额":
                replace_num = "{:.2f}".format(row[replace])
                docx_replace_regex(doc, regex, replace_num)
            elif replace == "今天日期":
                replace_date = row['今天日期']
                # replace_date = datetime.datetime.strptime(replace_str,"%Y-%m-%d %H:%M:%S") #字符串转日期
                replace_datef=f'{replace_date.year}年{replace_date.month}月{replace_date.day}日'
                docx_replace_regex(doc, regex, replace_datef)
            else:
                docx_replace_regex(doc, regex, row[replace])
            docx_replace_regex(doc.sections[0].header, regex, row[replace])

    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out','附件1评审结果说明-'+row["工程名称"]+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)

#写批复函数
def make_gov_approve_word(data):
    replacements = {
        r"评审工程名": "工程名称",
        r"送审报价": "送审金额",
        r"评审价": "审定金额",
        r"审减价": "审减金额",
        r"出报告时间": "出报告时间",
        r"送审甲方": "送审甲方",
        r"设计公司": "设计单位",
        r"送审预算单位": "送审预算编制单位",
        r"大写造价" : "大写造价",
        r"委托书编号" : "委托书编号",
        r"鑫诚报告号" : "鑫诚报告号",
        r"评审范围内容" : "评审范围内容",
        r"批复资金来源" : "批复资金来源",
    }      
    doc = Document(r'.\template_gushi\5批复模板.docx')  
    for index, row in data.iterrows():          
        for regex, replace in replacements.items():          
            regex = re.compile(regex, re.DOTALL)
            if replace ==  "送审金额"  or replace == "审定金额" or replace == "审减金额":
                replace_num = "{:.2f}".format(row[replace])
                docx_replace_regex(doc, regex, replace_num)
            elif replace == "出报告时间":
                replace_date = row[replace]
                # replace_date = datetime.datetime.strptime(replace_str,"%Y-%m-%d") #字符串转日期
                replace_datef=f'{replace_date.year}年{replace_date.month}月{replace_date.day}日'
                docx_replace_regex(doc, regex, replace_datef)
            else:
                docx_replace_regex(doc, regex, row[replace])
            
    # 将日期格式化为"年.月.日"的格式
    current_date = datetime.date.today()   
    formatted_date = current_date.strftime("%Y.%m.%d")
    value=str(data.iloc[0]['委托书编号'])[-4:-1] 
    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out', formatted_date+'-委'+value+'-关于'+row["工程名称"]+'预算的评审意见-鑫诚国际'+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)


#写报告函数
def make_report_word(data): 
    replacements = {
        r"评审工程名": "工程名称",
        r"送审报价": "送审金额",
        r"评审价": "审定金额",
        r"大写造价" : "大写造价",
        r"审减价": "审减金额",
        r"委托书开始时间": "委托书开始时间",
        r"出报告时间": "出报告时间",
        r"大写报告时间": "大写报告时间",
        r"送审甲方": "送审甲方",
        r"设计公司": "设计单位",
        r"送审预算单位": "送审预算编制单位",
        r"委托书编号" : "委托书编号",
        r"鑫诚报告号" : "鑫诚报告号",
        r"评审范围内容" : "评审范围内容",
        r"批复资金来源" : "批复资金来源",
    }      
    doc = Document(r'.\template_gushi\6审核报告模板.docx')  
    for index, row in data.iterrows():          
        for regex, replace in replacements.items():          
            regex = re.compile(regex, re.DOTALL)
            if replace ==  "送审金额"  or replace == "审定金额" or replace == "审减金额":
                replace_num = "{:.2f}".format(row[replace])
                docx_replace_regex(doc, regex, replace_num)
            elif replace == "出报告时间" or replace == "委托书开始时间":
                replace_date= row[replace]
                # replace_date = datetime.datetime.strptime(replace_str,"%Y-%m-%d") #字符串转日期
                replace_datef=f'{replace_date.year}年{replace_date.month}月{replace_date.day}日'
                docx_replace_regex(doc, regex, replace_datef)
            else:
                docx_replace_regex(doc, regex, row[replace])
            docx_replace_regex(doc.sections[0].header, regex, row[replace])
    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out', '2审核报告-'+row["工程名称"]+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)


#封面扉页
def make_reportface_word(data): 
    replacements = {
        r"评审工程名": "工程名称",
        r"出报告时间": "大写报告时间",
        r"鑫诚报告号" : "鑫诚报告号",
    }      
    doc1 = Document(r'.\template_gushi\6_1封面模板.docx') 
    doc2 = Document(r'.\template_gushi\6_2扉页模板.docx')   
    for index, row in data.iterrows():          
        for regex, replace in replacements.items():          
            regex = re.compile(regex, re.DOTALL)
            docx_replace_regex(doc1, regex, row[replace]) 
            docx_replace_regex(doc2, regex, row[replace])              

    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath1=os.path.join(current_dir,'out', '1.1封面-'+row["工程名称"]+'.docx')
    file_fullpath2=os.path.join(current_dir,'out', '1.2扉页-'+row["工程名称"]+'.docx')
    doc1.save(file_fullpath1)
    doc2.save(file_fullpath2)
    print('文件已经保存至'+file_fullpath1+'和'+file_fullpath2)
  

#编制说明
def make_explain_word(data):
    replacements = {
        r"评审工程名": "工程名称",
        r"评审范围内容" : "评审范围内容",
    }      
    doc = Document(r'.\template_gushi\7招标控制价总说明模板.docx')  
    for index, row in data.iterrows():          
        for regex, replace in replacements.items():          
            regex = re.compile(regex, re.DOTALL)
            docx_replace_regex(doc, regex, row[replace])
            docx_replace_regex(doc.sections[0].header, regex, row[replace])

    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out', '附件3招标控制价总说明-'+row["工程名称"]+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)

#写评审时间节点
def make_time_word(data):
    replacements = {
        r"评审工程名": "工程名称",
        r"委托书开始时间": "委托书开始时间",
        r"委托书结束时间": "委托书结束时间",
        r"出报告时间": "出报告时间",
        r"初审复核时间": "初审复核时间",
        r"上会的时间": "上会的时间",
        r"反馈的时间": "反馈的时间",
    }      
    doc = Document(r'.\template_gushi\评审时间节点模板.docx')  
    for index, row in data.iterrows():          
        for regex, replace in replacements.items():          
            regex = re.compile(regex, re.DOTALL)
            #用正则匹配replace所有含时间的关键词
            replace_time=re.findall(r'.*时间',replace)

            if replace_time :
                replace_date = row[replace]
                # replace_date = datetime.datetime.strptime(replace_str,"%Y-%m-%d") #字符串转日期
                replace_datef=f'{replace_date.year}年{replace_date.month}月{replace_date.day}日'
                docx_replace_regex(doc, regex, replace_datef)                
            else:
                docx_replace_regex(doc, regex, row[replace])
            docx_replace_regex(doc.sections[0].header, regex, row[replace])

    current_dir=os.getcwd()
    os.makedirs('out',exist_ok=True)
    file_fullpath=os.path.join(current_dir,'out', '评审时间节点-'+row["工程名称"]+'.docx')
    doc.save(file_fullpath)
    print('文件已经保存至'+file_fullpath)


#用户交互在定案表和初审复核表中选择
def main():
    while True:
        data=read_data()   
        while True:
            choice = input("1.初审复核表 2.审减明细 3.反馈函 4.定案表  5.批复、评审时间节点、处理签 6.报告及封皮 7.编制说明8.复核意见回复9.复核意见 \n输入0该项目退出,输入q退出程序 \n请输入你的选择:")
            if choice == "1":
                make_Application_word(data)  #1初审复核表
            elif choice == "8":
                make_check_reply_word(data)  #8复核意见回复
            elif choice == "9":
                make_check_word(data) #9复核意见
            elif choice == "2":
                make_deduction_details_word(data) #2审减明细
            elif choice == "3":
                make_feedback_word(data)   #3反馈函
            elif choice == "4":
                make_decision_sheet(data)   #4定案表
            elif choice == "5":
                make_gov_approve_word(data) #5批复
                make_govmark(data) #处理签
                make_time_word(data) #评审时间节点
            elif choice == "6":
                make_report_word(data)   #6报告
                make_reportface_word(data) #6报告封皮
            elif choice == "7":
                make_explain_word(data)  #7编制说明
            elif choice == "0":
                break             #退出
            elif choice == "q":
                exit()             #退出            
            else:
                choice = input("输入错误,重新输入:")

    
main()
        
